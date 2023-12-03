# -*- encoding:utf-8 -*-
import pandas as pd
import math
import os
import matplotlib.pyplot as plt # plt 用于显示图片
import matplotlib.image as mpimg # mpimg 用于读取图片
import jieba
from wordcloud import WordCloud
import csv
import matplotlib as  mpl
import matplotlib.pyplot as plt
import warnings
import seaborn as sns
import numpy as np
from sklearn.model_selection import train_test_split
from torch.utils.data import DataLoader, TensorDataset
from tqdm import tqdm
import torch.nn as nn
import torch
from keras.preprocessing.text import Tokenizer
from keras.preprocessing.sequence import pad_sequences
import pickle
import re
import json
import requests
from snownlp import SnowNLP
from collections import Counter
from snownlp import sentiment
import random
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
warnings.filterwarnings('ignore')
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from model import SentimentNet

# 用来显示中文标签
mpl.rcParams["font.family"] = "SimHei"
# 用来显示负号
mpl.rcParams["axes.unicode_minus"] = False

def extract_digits(string):
    pattern = r'\d+'
    matches = re.findall(pattern, string)
    digits = ''.join(matches)
    return digits

# 请求头定义
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/105.0.0.0 Safari/537.36"
}


# 获取课程参与人数和评论
def get_course_participate_comment(courseID):
    # session 实例化
    session = requests.session()
    # 慕课主页url
    index_url = "https://www.icourse163.org/"
    # 慕课主页发起请求，获取后面需要的csrfKey
    index_res = session.get(index_url, headers=HEADERS)
    # 获取csrfKey
    key = index_res.cookies.get("NTESSTUDYSI")

    # 详细课程url
    course_url = "https://www.icourse163.org/course/"+courseID+"?from=searchPage&outVendor=zw_mooc_pcssjg_"
    course_res = session.get(url=course_url,headers=HEADERS)
    # 获取课程参与人数
    deal = re.compile(r'enrollCount : "(.*?)"')
    result = deal.findall(course_res.text)
    participate_person = result[0]

    #课程评论url

    comment_url = f"https://www.icourse163.org/web/j/mocCourseV2RpcBean.getCourseEvaluatePaginationByCourseIdOrTermId.rpc?csrfKey={key}"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/105.0.0.0 Safari/537.36",
        "referer": "https://www.icourse163.org/course/"+courseID+"?from=searchPage&outVendor=zw_mooc_pcssjg_",
    }

    data_list = []


    with open("class_describe\\" + courseID + ".csv", mode="w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)

        # 获取所有评论数
        for i in range(1, 80):
            param = {
                "courseId": extract_digits(courseID),
                "pageIndex": i,
                "pageSize": "20",
                "orderBy": "3"
            }
            comment_res = session.post(url=comment_url, data=param, headers=headers)
            data = json.loads(comment_res.text)
            for count in range(len(data["result"]["list"])):
                agreecounts = data["result"]["list"][count]["agreeCount"]
                mark = data["result"]["list"][count]["mark"]
                content = data["result"]["list"][count]["content"]
                writer.writerow([content, agreecounts, mark])
                # 存储数据到列表
                data_list.append([content, agreecounts, mark])
            

    # 将列表转换为DataFrame
    df = pd.DataFrame(data_list, columns=["评论内容", "点赞数", "评分"])
    df = df.drop_duplicates()
    df = df.dropna()
    print("数据写入完毕！")
    return df

def calculate_likes_weight(likes):
    return (likes+1) / math.log2(likes + 2)

def calculate_likes_right(df):
        # 添加权值列
        df['点赞权值'] = df['点赞数'].apply(calculate_likes_weight)
        #评分情感
        df['评分情感'] = df['评分'].apply(lambda x: 1 if x > 3 else (-1 if x < 3 else 0))

def get_custom_stopword(stop_word_file):
        with open(stop_word_file, encoding='utf-8') as f:
            stop_word = f.read()

        stop_word_list = stop_word.split("\n")
        custom_stopword = [i for i in stop_word_list]
        return custom_stopword

def tokenized(df, stopwords):

    def tokenize_text(text):
        tokens = jieba.lcut(text)  # 使用结巴分词进行分词
        tokens_without_stopwords = [token for token in tokens if token not in stopwords]  # 去除停用词
        chinese_tokens = [token for token in tokens_without_stopwords if re.match(r'^[\u4e00-\u9fff]+$', token)]
        return chinese_tokens

    df['tokens'] = df['评论内容'].apply(tokenize_text)  # 对评论内容进行分词并保存到新的列中

    # 将分词结果展开为一个列表
    all_tokens = [token for tokens in df['tokens'] for token in tokens]

    # 计算词频
    word_frequencies = pd.Series(all_tokens).value_counts()

    # 获取出现频率最高的前五个词
    top_10_words = word_frequencies[:10].index.tolist()

    top_10_words_str = ', '.join(top_10_words)

    return top_10_words_str

def snownlp_predict(df):
        output_column = 'sentiment_prediction_snownlp'  # 新列的名称
        sentiment_column = 'sentiment_snownlp'  # 新列的名称

        # 将情感倾向预测结果保存到新列中
        df[output_column] = df['评论内容'].apply(lambda x: SnowNLP(x).sentiments)

        # 根据情感得分进行预测标签的转换
        df[sentiment_column] = df[output_column].apply(lambda x: 1 if x > 0.6 else (-1 if x < 0.4 else 0))




# 定义函数来进行分词和情感预测
def process_comments(comments):
    num_words = 3000
    model = SentimentNet(num_words, 256, 128, 8, 2)  # 创建一个新的模型实例
    model.load_state_dict(torch.load('model.pth'))

    sentence_len = 5

    # 进行分词和去除停用词
    comments_cut = [jieba.lcut(text) for text in comments]

    tokenizer_file = 'tokenizer.pkl'

    # 加载词汇表文件
    with open(tokenizer_file, 'rb') as f:
        tokenizer = pickle.load(f)


    # 将分词后的文本转换为序列
    test_seq = tokenizer.texts_to_sequences(texts=comments_cut)

    # 对序列进行填充
    test_pad_seq = pad_sequences(test_seq, maxlen=sentence_len, padding='post', truncating='post')

    # 初始化隐藏层
    h = model.init_hidden(len(test_pad_seq))

    # 获取情感倾向预测结果
    output, h = model(torch.tensor(test_pad_seq), h)

    # 返回情感倾向预测结果
    return output.view(-1, 2)

def LSTM_predict(df):
        output_column = 'sentiment_prediction_LSTM'  # 新列的名称
        sentiment_column = 'sentiment_LSTM'  # 新列的名称

        # 提取评论列
        comments = df['评论内容']

        # 处理评论并获取情感倾向预测结果
        predictions = process_comments(comments)

        # 将情感倾向预测结果保存到新列中
        df[output_column] = predictions.detach().cpu().numpy().tolist()

        # 添加"sentiment"列
        df[sentiment_column] = np.where(predictions.detach().cpu()[:, 0] > 0.6, 1, np.where(predictions.detach().cpu()[:, 0] < 0.4, -1, 0))

def remove_rows_by_value(df):
    # 删除指定列值为特定值的行
    df = df[df['check_sentiment'] != 2]

def compare_sentiments(df):
    # 将DataFrame转换为字典列表
    rows = df.to_dict('records')

    # 遍历每一行并逐行检查情感值
    for row in rows:
        # 获取当前行的情感值
        sentiment_value = row['评分情感']
        sentiment_snownlp_value = row['sentiment_snownlp']
        sentiment_LTSM_value = row['sentiment_LSTM']

        # 检查三列的内容是否一致
        if sentiment_value == sentiment_snownlp_value or sentiment_value == sentiment_LTSM_value:
            # 如果 sentiment_value 与 sentiment_snownlp_value 或 sentiment_LTSM_value 相同
            row['check_sentiment'] = sentiment_value
        elif sentiment_snownlp_value == sentiment_LTSM_value:
            # 如果 sentiment_snownlp_value 与 sentiment_LTSM_value 相同
            row['check_sentiment'] = sentiment_snownlp_value
        else:
            row['check_sentiment'] = 2

    # 将更新后的字典列表转换回DataFrame
    df = pd.DataFrame(rows)
    remove_rows_by_value(df)
    return df

    

def calculate_rank(df):

        # 计算含权评分
        df['标准含权评分'] = abs(df['评分'] - 3) * df['点赞权值'] * df['check_sentiment']

        # 计算∑权重*(评分-3)/评论数的结果
        status_weighted_rating = (df['标准含权评分'].sum()) / len(df)

        # 返回计算结果
        return status_weighted_rating



def get_infos(df,courseID):
    items = df['评论内容']
    item_lengths = items.str.len()
    length_counts = item_lengths.value_counts().sort_index()

    # 绘制分布图
    plt.bar(length_counts.index, length_counts.values)
    plt.xlabel('评论内容长度')
    plt.ylabel('Count')
    plt.title('Distribution of Item Lengths')
    plt.title('评论内容长度分布')
    plt.xlim(0, 100)
    plt.savefig('class_describe/'+courseID+'评论内容长度分布.png', bbox_inches='tight')
    plt.clf() 


    conditions = [df['点赞数'] == 0,
              df['点赞数'] == 1,
              df['点赞数'] == 2,
              df['点赞数'] >= 3]
    labels = ['点赞数为0', '点赞数为1', '点赞数为2', '点赞数为3以上']
    counts = [sum(condition) for condition in conditions]

    plt.figure(figsize=(6, 6))

    # 将counts的值添加到标签中
    labels_with_counts = [f'{label} ({count})' for label, count in zip(labels, counts)]

    plt.pie(counts, labels=labels_with_counts, autopct='%1.1f%%')

    plt.legend(loc='best')

    plt.savefig('class_describe/'+courseID+'点赞数分布情况.png')
    plt.clf() 

    
    labels = [f'{i}星' for i in range(1, 6)]
    sizes = [sum(np.ceil(df['评分']) == i) for i in range(1, 6)]
    colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#c2c2f0']

    plt.pie(sizes, labels=None, autopct='%1.1f%%', colors=colors)
    plt.title('评分的分布')
    # 显示图例
    plt.legend(labels, loc='best')
    # 保存图像
    plt.savefig('class_describe/'+courseID+'评分的分布.png')



    x = df['点赞数']
    y = np.ceil(df['评分']).astype(int) 

    # 根据点赞数和评分的值为数据点着色
    colors = df['点赞数']
    plt.clf() 
    plt.scatter(x, y, c=colors, cmap='viridis')
    plt.colorbar(label='点赞数')
    plt.xlabel('点赞数')
    plt.ylabel('评分')
    plt.title('点赞数与评分的分布的散点图')
    plt.yticks(np.arange(min(y), max(y)+1, 1))
    plt.savefig('class_describe/'+courseID+'点赞数与评分的分布的散点图.png')
    plt.clf() 

    #获取词云
    data = df['评论内容'].head(len(df['评论内容']) // 2)
    txt = ' '.join(data.tolist())
    string = ' '.join(jieba.lcut(txt))
    stopwords = get_custom_stopword("stopwords.txt")
    w = WordCloud(width=500,
                  height=400,
                  background_color='white',
                  font_path='也字工厂瑞云浓楷书.TTF',
                  scale=15,
                  stopwords=stopwords,
                  contour_width=5,
                  contour_color='red'
                  )

    w.generate(string)
    w.to_file('class_describe/'+courseID+'_wordcloud.png')

    #获取关键词
    top_10_words_str=tokenized(df,stopwords)
    data_df = pd.read_csv('new_all_course.csv')
    x = 10  # 初始关联词数量
    matching_rows = None
    while x >= 0 and (matching_rows is None or matching_rows.empty):
        matching_rows = data_df[(data_df['courseID'] != courseID) & (data_df['top_10_words'].astype(str).apply(lambda row: sum(1 for word in row.split(', ') if word.strip() in top_10_words_str.split(', '))) >= x)]
        x -= 1
    matching_rows_links_str = '\n'.join(matching_rows['link'].astype(str))
    
    calculate_likes_right(df)
    snownlp_predict(df)
    LSTM_predict(df)
    df=compare_sentiments(df)
    weighted_rating=calculate_rank(df)

    # 创建标准含权评分区间
    bin_edges = np.arange(-0.5, 7, 0.5)  # 设置评分区间边界

    alldata_df = pd.read_csv('new_all_course.csv')
    # 统计标准含权评分在各区间的数量
    scores = alldata_df['标准含权评分']
    score_counts, _ = pd.cut(scores, bins=bin_edges, include_lowest=True, right=False, retbins=True)
    score_counts = score_counts.value_counts().sort_index()
    # 绘制金字塔图
    plt.clf() 
    plt.figure(figsize=(8, 6))
    # 绘制条形图
    bars = plt.barh(score_counts.index.astype(str)[::-1], score_counts[::-1], color='lightblue')
    # 标记输入值所在的区间
    input_value = weighted_rating  # 替换为你的输入值
    # 显示每个区间的数量
    for i, bar in enumerate(bars):
        plt.text(40+bar.get_width(), i, str(int(bar.get_width())), va='center', ha='right')
    plt.xlabel('数量')
    plt.ylabel('标准含权评分区间')
    plt.title('标准含权评分分布')
    # 给出输入值在标准含权评分的排名
    rank = (scores >= input_value).sum()
    # 显示图形
    plt.savefig('class_describe/'+courseID+'标准含权评分分布.png')

    doc = Document()

    doc.add_picture('class_describe/'+courseID+'评论内容长度分布.png', width=Inches(6), height=Inches(4))
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1.add_run("评论内容长度分布如图").bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_picture('class_describe/'+courseID+'点赞数分布情况.png', width=Inches(6), height=Inches(6))
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.add_run("点赞数分布情况如图").bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_picture('class_describe/'+courseID+'评分的分布.png', width=Inches(6), height=Inches(6))
    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p3.add_run("评分的分布情况如图").bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_picture('class_describe/'+courseID+'点赞数与评分的分布的散点图.png', width=Inches(6), height=Inches(6))
    p4 = doc.add_paragraph()
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p4.add_run("点赞数与评分的分布的散点图如图").bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_picture('class_describe/'+courseID+'_wordcloud.png', width=Inches(6), height=Inches(6))
    p5 = doc.add_paragraph()
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p5.add_run(courseID+"评论词云如图").bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    doc.add_picture('class_describe/'+courseID+'标准含权评分分布.png', width=Inches(8), height=Inches(6))
    p7 = doc.add_paragraph()
    p7.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p7.add_run(courseID+"标准含权评分分布如图").bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


    p6 = doc.add_paragraph()
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p6.add_run(courseID + "评论的十大关键词:" + top_10_words_str + '\n\n\n本课程在所有课程的质量的排名：' + str(rank))
    run.font.size = Pt(20)

    p8 = doc.add_paragraph()
    p8.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p8.add_run('\n\n\n'+courseID + "可能相关的课程:" + matching_rows_links_str )
    run.font.size = Pt(20)


    # 保存文档
    doc.save('class_describe/describe_'+courseID+'.docx')

    
    

if __name__ == '__main__':
    link=input("请输入课程主页网址:")
    courseID = link.split('/')[-1]
    df=get_course_participate_comment(courseID)
    get_infos(df,courseID)
    df.to_csv('class_describe/describe_'+courseID+'.csv', index=False)